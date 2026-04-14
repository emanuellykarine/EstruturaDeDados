from django.http import HttpResponse
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import cm
from datetime import date
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO

def formatar_data(data):
    if isinstance(data, date):
        return data.strftime('%d/%m/%Y')
    if isinstance(data, str) and data:
        try:
            # Tenta converter string ISO (YYYY-MM-DD) para dd/mm/yyyy
            partes = data.split('-')
            if len(partes) == 3:
                return f"{partes[2]}/{partes[1]}/{partes[0]}"
        except:
            pass
    return str(data) if data else ''

def get_status_texto(status):
    return 'Ativo' if status in [True, 'ativo', 'Ativo'] else 'Inativo'

def gerar_docx(titulo, dados):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="relatorio.docx"'

    document = Document()
    
    # Título
    titulo_paragraph = document.add_heading(titulo, level=1)
    titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cabeçalhos da tabela
    cabecalho = ['Aluno', 'Status', 'Tipo Auxílio', 'Curso', 'Data Início', 'Data Fim']
    
    # Criar tabela
    table = document.add_table(rows=1, cols=len(cabecalho))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Estilo do cabeçalho
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(cabecalho):
        cell = header_cells[i]
        cell.text = header_text
        # Cor de fundo azul
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        # Texto branco e negrito
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = None  # Branco não funciona bem, então usamos negrito
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adicionar dados
    for item in dados:
        aluno_nome = str(item.get('aluno', {}).get('full_name', '') if isinstance(item.get('aluno'), dict) else item.get('aluno', ''))
        curso_nome = str(item.get('curso', ''))
        tipo_auxilio = str(item.get('tipo_auxilio', ''))
        
        row_cells = table.add_row().cells
        row_data = [
            aluno_nome,
            get_status_texto(item.get('status')),
            tipo_auxilio,
            curso_nome,
            formatar_data(item.get('data_inicio', '')),
            formatar_data(item.get('data_fim', ''))
        ]
        
        for i, value in enumerate(row_data):
            row_cells[i].text = value
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # Ajustar largura das colunas
    widths = [Cm(4), Cm(2), Cm(3.5), Cm(7), Cm(2.5), Cm(2.5)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Salvar no response
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    response.write(buffer.getvalue())
    
    return response

def gerar_pdf(titulo, dados):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="relatorio.pdf"'

    doc = SimpleDocTemplate(
        response,
        pagesize=landscape(A4),
        rightMargin=1*cm,
        leftMargin=1*cm,
        topMargin=1*cm,
        bottomMargin=1*cm
    )

    elements = []
    styles = getSampleStyleSheet()

    # Título
    titulo_style = styles['Heading1']
    titulo_style.alignment = 1  # Centralizado
    elements.append(Paragraph(titulo, titulo_style))
    elements.append(Spacer(1, 0.5*cm))

    # Estilo para células com quebra de linha
    cell_style = styles['Normal']
    cell_style.fontSize = 9
    cell_style.leading = 11

    # Cabeçalho da tabela
    cabecalho = ['Aluno', 'Status', 'Tipo Auxílio', 'Curso', 'Data Início', 'Data Fim']

    # Dados da tabela
    table_data = [cabecalho]
    for item in dados:
        aluno_nome = str(item.get('aluno', '').get('full_name', '') if isinstance(item.get('aluno'), dict) else item.get('aluno', ''))
        curso_nome = str(item.get('curso', ''))
        tipo_auxilio = str(item.get('tipo_auxilio', ''))
        
        row = [
            Paragraph(aluno_nome, cell_style),
            'Ativo' if item.get('status') in [True, 'ativo', 'Ativo'] else 'Inativo',
            Paragraph(tipo_auxilio, cell_style),
            Paragraph(curso_nome, cell_style),
            formatar_data(item.get('data_inicio', '')),
            formatar_data(item.get('data_fim', ''))
        ]
        table_data.append(row)

    # Criar tabela
    col_widths = [4.5*cm, 2*cm, 4*cm, 8*cm, 2.5*cm, 2.5*cm]
    table = Table(table_data, colWidths=col_widths, repeatRows=1)

    # Estilo da tabela (visual de planilha)
    table_style = TableStyle([
        # Cabeçalho
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 8),

        # Corpo da tabela
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('TOPPADDING', (0, 1), (-1, -1), 6),

        # Bordas (estilo planilha)
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#B4C6E7')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#4472C4')),

        # Linhas alternadas (zebra)
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#D6DCE5')]),
    ])

    table.setStyle(table_style)
    elements.append(table)

    doc.build(elements)
    return response


def gerar_xlsx(titulo, dados):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="relatorio.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Relatório"

    # Estilos
    header_font = Font(bold=True, color="FFFFFF", size=10)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    
    cell_font = Font(size=9)
    cell_alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    
    thin_border = Border(
        left=Side(style='thin', color='B4C6E7'),
        right=Side(style='thin', color='B4C6E7'),
        top=Side(style='thin', color='B4C6E7'),
        bottom=Side(style='thin', color='B4C6E7')
    )
    
    zebra_fill = PatternFill(start_color="D6DCE5", end_color="D6DCE5", fill_type="solid")

    # Título
    ws.merge_cells('A1:F1')
    ws['A1'] = titulo
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    # Cabeçalho
    cabecalho = ['Aluno', 'Status', 'Tipo Auxílio', 'Curso', 'Data Início', 'Data Fim']
    for col, header in enumerate(cabecalho, 1):
        cell = ws.cell(row=3, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Dados
    for row_idx, item in enumerate(dados, 4):
        aluno_nome = str(item.get('aluno', {}).get('full_name', '') if isinstance(item.get('aluno'), dict) else item.get('aluno', ''))
        curso_nome = str(item.get('curso', ''))
        tipo_auxilio = str(item.get('tipo_auxilio', ''))
        
        row_data = [
            aluno_nome,
            get_status_texto(item.get('status')),
            tipo_auxilio,
            curso_nome,
            formatar_data(item.get('data_inicio', '')),
            formatar_data(item.get('data_fim', ''))
        ]
        
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = cell_font
            cell.alignment = cell_alignment
            cell.border = thin_border
            
            # Linhas alternadas (zebra)
            if (row_idx - 4) % 2 == 1:
                cell.fill = zebra_fill

    # Ajustar largura das colunas
    col_widths = [30, 12, 25, 50, 15, 15]
    for col_idx, width in enumerate(col_widths, 1):
        ws.column_dimensions[chr(64 + col_idx)].width = width

    # Salvar no response
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    response.write(buffer.getvalue())
    
    return response

